from pathlib import Path
import shutil

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


TITLE = "\u9898\u5e93\u5bfc\u5165\u6807\u51c6\u6a21\u677f"
FILE_NAME_ASCII = "question-bank-import-template.docx"
FILE_NAME_CN = "\u9898\u5e93\u5bfc\u5165\u6807\u51c6\u6a21\u677f.docx"

HEADERS = [
    "\u9898\u76ee",
    "\u9898\u578b",
    "\u80fd\u529b\u70b9",
    "\u96be\u5ea6",
    "\u53c2\u8003\u7b54\u6848",
    "\u8bc4\u5206\u89c4\u5219",
    "\u8ffd\u95ee\u9898",
    "\u6765\u6e90",
]

ROWS = [
    [
        "\u4f60\u4e3a\u4ec0\u4e48\u62a5\u6211\u4eec\u5b66\u6821\u7684\u4eba\u5de5\u667a\u80fd\u6280\u672f\u5e94\u7528\u8fd9\u4e2a\u4e13\u4e1a\u5462\uff1f",
        "\u7efc\u5408\u9898",
        "\u4e13\u4e1a\u8ba4\u77e5",
        "medium",
        "\u6211\u62a5\u8003\u8fd9\u4e2a\u4e13\u4e1a\uff0c\u4e3b\u8981\u662f\u56e0\u4e3a\u6211\u5bf9\u4eba\u5de5\u667a\u80fd\u9886\u57df\u6709\u6301\u7eed\u5174\u8da3\uff0c\u4e5f\u770b\u5230\u4e86\u5b83\u5728\u4ea7\u4e1a\u4e2d\u7684\u5e7f\u6cdb\u5e94\u7528\u3002\u6211\u5e0c\u671b\u5728\u5927\u5b66\u9636\u6bb5\u7cfb\u7edf\u5b66\u4e60\u7f16\u7a0b\u3001\u6570\u636e\u5206\u6790\u3001\u673a\u5668\u5b66\u4e60\u548c\u5b9e\u9645\u9879\u76ee\u5f00\u53d1\u80fd\u529b\uff0c\u4e3a\u4ee5\u540e\u4ece\u4e8b\u667a\u80fd\u5e94\u7528\u5f00\u53d1\u6253\u4e0b\u57fa\u7840\u3002",
        "\u4e13\u4e1a\u8ba4\u77e540%\uff1b\u8868\u8fbe\u903b\u8f9130%\uff1b\u804c\u4e1a\u89c4\u521230%",
        "\u4f60\u4e86\u89e3\u8fd9\u4e2a\u4e13\u4e1a\u672a\u6765\u4e3b\u8981\u4f1a\u5b66\u4e60\u54ea\u4e9b\u5185\u5bb9\u5417\uff1f",
        "\u6821\u5185\u6574\u7406",
    ],
    [
        "\u8fd9\u4e2a\u4e13\u4e1a\u4ee5\u540e\u80fd\u505a\u4ec0\u4e48\uff1f",
        "\u7efc\u5408\u9898",
        "\u804c\u4e1a\u89c4\u5212",
        "medium",
        "\u8fd9\u4e2a\u4e13\u4e1a\u6bd5\u4e1a\u540e\u53ef\u4ee5\u4ece\u4e8b\u4eba\u5de5\u667a\u80fd\u5e94\u7528\u5f00\u53d1\u3001\u6570\u636e\u5904\u7406\u3001\u7b97\u6cd5\u8f85\u52a9\u5f00\u53d1\u3001\u667a\u80fd\u4ea7\u54c1\u6d4b\u8bd5\u4e0e\u5b9e\u65bd\u7b49\u5de5\u4f5c\uff0c\u4e5f\u53ef\u4ee5\u7ee7\u7eed\u6df1\u9020\uff0c\u5f80\u7b97\u6cd5\u3001\u6570\u636e\u5206\u6790\u6216\u8f6f\u4ef6\u5f00\u53d1\u65b9\u5411\u53d1\u5c55\u3002",
        "\u5c97\u4f4d\u7406\u89e340%\uff1b\u4e3e\u4f8b\u5177\u4f5330%\uff1b\u8868\u8fbe\u5b8c\u657430%",
        "\u5982\u679c\u4f60\u8fdb\u5165\u8fd9\u4e2a\u4e13\u4e1a\uff0c\u6700\u60f3\u5c1d\u8bd5\u54ea\u4e2a\u5c31\u4e1a\u65b9\u5411\uff1f",
        "\u6821\u5185\u6574\u7406",
    ],
    [
        "\u8bf7\u4ecb\u7ecd\u4e00\u4e0b\u4f60\u7406\u89e3\u7684\u673a\u5668\u5b66\u4e60\u3002",
        "\u4e13\u4e1a\u9898",
        "\u4e13\u4e1a\u77e5\u8bc6",
        "medium",
        "\u6211\u7406\u89e3\u7684\u673a\u5668\u5b66\u4e60\uff0c\u5c31\u662f\u8ba9\u8ba1\u7b97\u673a\u901a\u8fc7\u6570\u636e\u5b66\u4e60\u89c4\u5f8b\uff0c\u518d\u5229\u7528\u8fd9\u4e9b\u89c4\u5f8b\u5b8c\u6210\u9884\u6d4b\u6216\u5206\u7c7b\u4efb\u52a1\u3002\u5b83\u901a\u5e38\u5305\u62ec\u6570\u636e\u51c6\u5907\u3001\u7279\u5f81\u5904\u7406\u3001\u6a21\u578b\u8bad\u7ec3\u3001\u6548\u679c\u8bc4\u4f30\u548c\u6301\u7eed\u4f18\u5316\u51e0\u4e2a\u6b65\u9aa4\u3002",
        "\u6982\u5ff5\u51c6\u786e40%\uff1b\u7ed3\u6784\u6e05\u667030%\uff1b\u4e3e\u4f8b\u8bf4\u660e30%",
        "\u4f60\u77e5\u9053\u76d1\u7763\u5b66\u4e60\u548c\u975e\u76d1\u7763\u5b66\u4e60\u6709\u4ec0\u4e48\u533a\u522b\u5417\uff1f",
        "\u6559\u5e08\u81ea\u7f16",
    ],
]


def set_east_asia_font(run, font_name: str) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def build_document() -> Document:
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    normal = doc.styles["Normal"]
    normal.font.size = Pt(10.5)
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run(TITLE)
    title_run.bold = True
    title_run.font.size = Pt(16)
    set_east_asia_font(title_run, "Microsoft YaHei")

    p1 = doc.add_paragraph()
    p1.add_run("\u6587\u4ef6\u683c\u5f0f\uff1a").bold = True
    p1.add_run(".docx\uff08Word \u6587\u6863\uff09")

    p2 = doc.add_paragraph()
    p2.add_run("\u586b\u5199\u8981\u6c42\uff1a").bold = True
    p2.add_run(
        "\u4f7f\u7528\u4e00\u5f20\u89c4\u8303\u8868\u683c\uff1b\u6bcf\u4e00\u884c\u4ee3\u8868\u4e00\u9053\u9898\uff1b"
        "\u4e0d\u8981\u5408\u5e76\u5355\u5143\u683c\uff1b\u4e0d\u8981\u6539\u52a8\u8868\u5934\u540d\u79f0\u3002"
    )

    table = doc.add_table(rows=1, cols=len(HEADERS))
    table.style = "Table Grid"
    for index, text in enumerate(HEADERS):
        table.rows[0].cells[index].text = text

    for row in ROWS:
        cells = table.add_row().cells
        for index, text in enumerate(row):
            cells[index].text = text

    p3 = doc.add_paragraph()
    p3.add_run("\u53ef\u9009\u503c\u5efa\u8bae\uff1a").bold = True
    p3.add_run(
        "\u9898\u578b\u53ef\u7528\u201c\u7efc\u5408\u9898 / \u4e13\u4e1a\u9898 / \u7ed3\u6784\u5316\u9898\u201d\uff1b"
        "\u96be\u5ea6\u53ef\u7528\u201ceasy / medium / hard\u201d\u3002"
    )

    return doc


def main() -> None:
    output_dir = Path("docs/templates")
    output_dir.mkdir(parents=True, exist_ok=True)

    ascii_path = output_dir / FILE_NAME_ASCII
    cn_path = output_dir / FILE_NAME_CN

    document = build_document()
    document.save(ascii_path)
    shutil.copyfile(ascii_path, cn_path)

    print(ascii_path.resolve())
    print(cn_path.resolve())


if __name__ == "__main__":
    main()
